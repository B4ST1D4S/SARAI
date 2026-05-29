#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Genera PROPUESTA_COMERCIAL_SARAI.docx — version profesional"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\proyect\EstetIA\docs\PROPUESTA_COMERCIAL_SARAI.docx"

# ── Paleta de colores ─────────────────────────────────────────────────────────
def rgb(h): return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

GOLD="C9A84C"; GDARK="8B5E0A"; GBG="FFFBEF"; GBORDER="E8C97A"
NAVY="1A2744"; NAVY2="2D3F6B"; WHITE="FFFFFF"
LGRAY="F4F4F4"; MGRAY="AAAAAA"; DGRAY="444444"
EMERALD="047857"; ROSE="B91C1C"; CYAN="0369A1"; VIOLET="5B21B6"
CHECK="✓"; CROSS="—"; SOON="Proximamente"

# ── Helpers XML ───────────────────────────────────────────────────────────────
def cell_bg(cell, color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), color); tcPr.append(shd)

def cell_valign(cell, val='center'):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    v = OxmlElement('w:vAlign'); v.set(qn('w:val'), val); tcPr.append(v)

def tbl_borders(table, color="C9A84C", sz=4):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    bd = OxmlElement('w:tblBorders')
    for n in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{n}')
        b.set(qn('w:val'),'single'); b.set(qn('w:sz'),str(sz))
        b.set(qn('w:space'),'0'); b.set(qn('w:color'),color)
        bd.append(b)
    tblPr.append(bd)

def no_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr'); tbl.insert(0, tblPr)
    bd = OxmlElement('w:tblBorders')
    for n in ['top','left','bottom','right','insideH','insideV']:
        b = OxmlElement(f'w:{n}')
        b.set(qn('w:val'),'none'); b.set(qn('w:sz'),'0')
        b.set(qn('w:space'),'0'); b.set(qn('w:color'),'auto')
        bd.append(b)
    tblPr.append(bd)

# ── Helpers contenido ─────────────────────────────────────────────────────────
def rn(para, text, bold=False, color=DGRAY, size=10.5, italic=False):
    r = para.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.color.rgb = rgb(color); r.font.size = Pt(size)
    return r

def para_bg(para, color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), color); pPr.append(shd)

def section_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(28); p.paragraph_format.space_after = Pt(4)
    rn(p, "  " + text.upper(), bold=True, color=GOLD, size=8.5)

def heading(doc, text, hl=None, size=18):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(10)
    if hl and hl in text:
        parts = text.split(hl, 1)
        rn(p, parts[0], bold=True, color=NAVY, size=size)
        rn(p, hl, bold=True, color=GOLD, size=size)
        if len(parts) > 1: rn(p, parts[1], bold=True, color=NAVY, size=size)
    else:
        rn(p, text, bold=True, color=NAVY, size=size)

def body(doc, text, size=10.5):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    rn(p, text, color=DGRAY, size=size)

def divider(doc, color=GBORDER):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6')
    bot.set(qn('w:space'),'1'); bot.set(qn('w:color'), color)
    pBdr.append(bot); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)

def tbl_header_row(table, row_idx, headers, bgs, text_color=WHITE, size=9.5):
    for col_idx, (hdr, bg) in enumerate(zip(headers, bgs)):
        c = table.cell(row_idx, col_idx)
        cell_bg(c, bg)
        cell_valign(c)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn(p, hdr, bold=True, color=text_color, size=size)

# ════════════════════════════════════════════════════════════════════════════════
# CREAR DOCUMENTO
# ════════════════════════════════════════════════════════════════════════════════
doc = Document()

for sec in doc.sections:
    sec.page_width  = Cm(21); sec.page_height = Cm(29.7)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.5)
    sec.top_margin  = Cm(2.5); sec.bottom_margin = Cm(2.0)

# Estilo base
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)

# ── PORTADA ───────────────────────────────────────────────────────────────────
badge = doc.add_paragraph()
badge.alignment = WD_ALIGN_PARAGRAPH.CENTER
badge.paragraph_format.space_before = Pt(36); badge.paragraph_format.space_after = Pt(20)
rn(badge, "SISTEMA PREMIUM  ·  INTELIGENCIA ARTIFICIAL MEDICA  ·  COLOMBIA 2026",
   bold=True, color=GOLD, size=8.5)

# Titulo SARAI
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_before = Pt(0); t.paragraph_format.space_after = Pt(4)
rn(t, "SARAI", bold=True, color=GOLD, size=72)

sub1 = doc.add_paragraph()
sub1.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub1.paragraph_format.space_before = Pt(0); sub1.paragraph_format.space_after = Pt(4)
rn(sub1, "Sistema de Asistencia, Registro e Automatizacion con Inteligencia Artificial",
   bold=True, color=NAVY, size=13)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub2.paragraph_format.space_before = Pt(0); sub2.paragraph_format.space_after = Pt(20)
rn(sub2, "Plataforma Integral de Gestion Clinica Impulsada por IA para Colombia",
   color=MGRAY, size=11)

divider(doc)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.paragraph_format.space_after = Pt(28)
rn(meta, "Oferta Comercial Confidencial  |  Version 2.0  |  Mayo 2026", color=MGRAY, size=9)

# Tabla de stats portada
ct = doc.add_table(rows=2, cols=5); ct.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_borders(ct, color=GBORDER, sz=6)
stats = [("+20","Modulos activos"),("IA","Asistente de voz"),("100%","Cloud / On-Premise"),("RIPS","Normativa CO"),("24/7","Disponibilidad")]
for ci, (val, lbl) in enumerate(stats):
    for ri, (txt, c, sz, bld) in enumerate([(val, GDARK, 20, True),(lbl, MGRAY, 8, False)]):
        cell = ct.cell(ri, ci)
        cell_bg(cell, GBG)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rn(p, txt, bold=bld, color=c, size=sz)
    ct.columns[ci].width = Cm(3.4)

doc.add_paragraph()
conf = doc.add_paragraph()
conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn(conf, "SARAI Technologies  |  comercial@sarai.com.co  |  Colombia 2026",
   color=GDARK, size=9, italic=True)

doc.add_page_break()

# ── 01 RESUMEN EJECUTIVO ──────────────────────────────────────────────────────
section_label(doc, "01  Resumen Ejecutivo")
heading(doc, "Una plataforma que transforma la clinica moderna", "transforma")
body(doc, "SARAI es el primer sistema colombiano de gestion clinica integral potenciado con Inteligencia Artificial conversacional. Automatiza flujos operativos, elimina digitacion manual, garantiza cumplimiento normativo y convierte cada consulta en una experiencia medica de clase mundial.")

# Stats table
st = doc.add_table(rows=5, cols=2)
tbl_borders(st, color=GBORDER, sz=4)
stats_rows = [
    ("80%", "Reduccion de tiempo en registro clinico con dictado por voz"),
    ("0 Errores", "Con IA de transcripcion — cero digitacion manual"),
    ("100%", "Conformidad RIPS / Resolucion 1995/1999 Ministerio de Salud"),
    ("24/7", "Disponibilidad cloud garantizada mediante contrato SLA"),
    ("Infinita", "Escalabilidad sin limite de usuarios ni sedes"),
]
for ri, (val, desc) in enumerate(stats_rows):
    vc = st.cell(ri, 0); dc = st.cell(ri, 1)
    bg_v = NAVY if ri % 2 == 0 else NAVY2
    cell_bg(vc, bg_v); cell_bg(dc, LGRAY if ri % 2 == 0 else WHITE)
    cell_valign(vc); cell_valign(dc)
    pv = vc.paragraphs[0]; pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(pv, val, bold=True, color=GOLD, size=16)
    pd = dc.paragraphs[0]
    rn(pd, desc, color=NAVY, size=10.5)
    vc.width = Cm(4); dc.width = Cm(12)

doc.add_paragraph()
# Problema vs Solucion
ps = doc.add_table(rows=1, cols=2); tbl_borders(ps, color=GBORDER, sz=4)
lc = ps.cell(0,0); rc = ps.cell(0,1)
cell_bg(lc, "FFF0F0"); cell_bg(rc, "F0FFF5")
rn(lc.paragraphs[0], "SIN SARAI — Situacion Actual", bold=True, color=ROSE, size=11)
for t in ["3-5 horas diarias en papeleo y digitacion",
          "Historias clinicas incompletas o extraviadas",
          "Confirmaciones de cita manuales por telefono",
          "Errores en RIPS = glosas y devoluciones",
          "Busqueda de pacientes en Excel o papel",
          "Sin metricas ni reportes en tiempo real",
          "Riesgo de incumplimiento normativo"]:
    p = lc.add_paragraph("  x  " + t); p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(9.5); p.runs[0].font.color.rgb = rgb(DGRAY)

rn(rc.paragraphs[0], "CON SARAI — Transformacion Digital", bold=True, color=EMERALD, size=11)
for t in ["Historia clinica completa en menos de 2 minutos",
          "Dictado por voz — cero digitacion",
          "Correos y confirmaciones 100% automaticos",
          "RIPS sin errores, generados automaticamente",
          "Busqueda de paciente en menos de 3 segundos",
          "Dashboard con metricas en tiempo real",
          "100% alineado con normativa colombiana"]:
    p = rc.add_paragraph("  v  " + t); p.paragraph_format.space_after = Pt(2)
    p.runs[0].font.size = Pt(9.5); p.runs[0].font.color.rgb = rgb(DGRAY)

lc.width = Cm(8); rc.width = Cm(8)
doc.add_page_break()

# ── 02 INTELIGENCIA ARTIFICIAL ────────────────────────────────────────────────
section_label(doc, "02  Inteligencia Artificial")
heading(doc, "El corazon inteligente del sistema", "inteligente")
body(doc, "SARAI va mucho mas alla de un HIS tradicional. Su asistente de IA escucha, entiende y actua sobre la informacion clinica en tiempo real.")

ai = [
    ("DICTADO POR VOZ — Whisper AI",
     "El medico dicta la historia clinica en tiempo real. Whisper Medium convierte voz a texto con 95%+ de precision en espanol medico colombiano, eliminando la digitacion completamente."),
    ("ASISTENTE SARAI CHAT",
     "Asistente conversacional que responde preguntas clinicas, sugiere diagnosticos CIE-10, ayuda con protocolos y genera resumenes automaticos de la consulta."),
    ("ANALITICA PREDICTIVA",
     "Paneles en tiempo real con metricas de ocupacion, productividad medica, cancelaciones e indicadores financieros para decisiones basadas en datos."),
    ("AUTOMATIZACION DE FLUJOS",
     "Correos de confirmacion, recordatorios, cotizaciones y documentos se disparan automaticamente sin intervencion humana."),
    ("SEGURIDAD EMPRESARIAL",
     "JWT con rotacion automatica, cifrado AES-256 en reposo y TLS en transito. Auditoria completa conforme a Ley 1581/2012 (habeas data Colombia)."),
    ("MULTI-SEDE / MULTI-EMPRESA",
     "Arquitectura multi-tenant que soporta multiples sedes y clinicas desde un unico panel de administracion centralizado."),
]
ai_t = doc.add_table(rows=3, cols=2); tbl_borders(ai_t, color=GBORDER, sz=4)
for i, (title, desc) in enumerate(ai):
    ri = i // 2; ci = i % 2
    c = ai_t.cell(ri, ci)
    cell_bg(c, GBG if ci == 0 else WHITE)
    rn(c.paragraphs[0], title, bold=True, color=NAVY, size=10)
    p2 = c.add_paragraph(); rn(p2, desc, color=DGRAY, size=9.5)
    c.width = Cm(8)

# ── 03 MODULOS ────────────────────────────────────────────────────────────────
section_label(doc, "03  Modulos del Sistema")
heading(doc, "Cobertura total del ciclo clinico", "total")
body(doc, "Desde la primera cita hasta la facturacion RIPS — cada proceso conectado y automatizado en una sola plataforma.")

mods = [
    "Agenda Inteligente Ilimitada","Historia Clinica Digital",
    "Gestion Integral de Pacientes","Cotizaciones Automaticas",
    "Facturacion RIPS","Consulta Externa",
    "Cirugia y Procedimientos","Hospitalizacion",
    "CUPS (Res. 5521/2013)","Configuracion de Disponibilidad",
    "Parametrizacion del Sistema","Gestion de Usuarios y Roles",
    "Central de Impresion","Panel de Administracion",
    "Asistente IA (SARAI)","Interoperabilidad HL7 (Proximo)",
]
COLS = 4; rows = (len(mods) + COLS - 1) // COLS
mt = doc.add_table(rows=rows, cols=COLS); tbl_borders(mt, color=GBORDER, sz=4)
for idx, name in enumerate(mods):
    ri = idx // COLS; ci = idx % COLS
    c = mt.cell(ri, ci)
    cell_bg(c, GBG if (ri+ci) % 2 == 0 else WHITE)
    rn(c.paragraphs[0], name, bold=True, color=NAVY, size=9.5)
    c.width = Cm(4)
doc.add_page_break()

# ── 04 PAQUETES ───────────────────────────────────────────────────────────────
section_label(doc, "04  Paquetes de Servicio")
heading(doc, "Elige el paquete que escala con tu clinica", "escala")
body(doc, "Cada paquete incluye onboarding, capacitacion, soporte tecnico y actualizaciones. Precios en COP + IVA mensual por sede.")

# ─ Tabla resumen 3 paquetes ───────────────────────────────────────────────────
p_s = doc.add_table(rows=8, cols=4)
tbl_borders(p_s, color="C9A84C", sz=6)
p_s.alignment = WD_TABLE_ALIGNMENT.CENTER

# Cabeceras
tbl_header_row(p_s, 0,
    ["CARACTERISTICA", "AMBULATORIO", "CIRUGIA & ESTETICA  *  MAS POPULAR", "HOSPITALIZACION"],
    [NAVY, CYAN, GDARK, VIOLET], size=9)

rows_data = [
    ("Precio mensual + IVA",   "$1.490.000",  "$3.290.000",  "$5.990.000"),
    ("Implementacion (unico)", "$2.500.000",  "$5.800.000",  "$12.000.000"),
    ("Usuarios concurrentes",  "Hasta 3",     "Hasta 15",    "Ilimitados"),
    ("Almacenamiento cloud",   "10 GB",       "50 GB",       "500 GB"),
    ("Soporte tecnico",        "Ticket 5x8",  "Prioritario 7x12","24/7 con SLA"),
    ("RIPS automatico",        "No incluido", "Incluido",    "Incluido"),
    ("Cirugia / Hospitalizacion","No incluido","Cirugia",    "Ambos"),
]
alt = [WHITE, LGRAY, WHITE, LGRAY, WHITE, LGRAY, WHITE]
for ri, (feat, a, b, c) in enumerate(rows_data):
    row_i = ri + 1
    vals = [feat, a, b, c]
    bgs  = ["F8F4E8", alt[ri], GBG, alt[ri]]
    colors = [NAVY, DGRAY, GDARK, DGRAY]
    bolds  = [True, False, True, False]
    for ci, (v, bg, co, bd) in enumerate(zip(vals, bgs, colors, bolds)):
        cell = p_s.cell(row_i, ci)
        cell_bg(cell, bg); cell_valign(cell)
        pp = cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
        rn(pp, v, bold=bd, color=co, size=10)

col_ws = [Cm(5.5), Cm(3.8), Cm(4.2), Cm(4.2)]
for row in p_s.rows:
    for ci, w in enumerate(col_ws): row.cells[ci].width = w

doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn(note, "(*) El paquete Cirugia & Estetica es el mas recomendado para clinicas de medicina estetica y procedimientos especializados.",
   color=GDARK, size=9, italic=True)
doc.add_paragraph()

# ─ Tabla de caracteristicas por paquete ──────────────────────────────────────
ht = doc.add_paragraph(); rn(ht, "MODULOS INCLUIDOS POR PAQUETE", bold=True, color=NAVY, size=11)
ht.paragraph_format.space_before = Pt(4); ht.paragraph_format.space_after = Pt(8)

feat_t = doc.add_table(rows=1, cols=4); tbl_borders(feat_t, color="C9A84C", sz=4)
tbl_header_row(feat_t, 0, ["MODULO / FUNCION", "AMBULATORIO", "CIRUGIA", "HOSPITALIZACION"],
               [NAVY, NAVY, NAVY, NAVY], size=9)

sections_feats = [
    ("CONSULTA EXTERNA", None, None, None),
    ("Agenda ilimitada",           CHECK,CHECK,CHECK),
    ("Historia clinica digital",   CHECK,CHECK,CHECK),
    ("Dictado voz (Whisper AI)",   CHECK,CHECK,CHECK),
    ("Cotizaciones automaticas",   CHECK,CHECK,CHECK),
    ("Asistente IA SARAI",         CHECK,CHECK,CHECK),
    ("CIRUGIA Y PROCEDIMIENTOS",   None,None,None),
    ("Modulo de cirugia",          CROSS,CHECK,CHECK),
    ("Consentimientos digitales",  CROSS,CHECK,CHECK),
    ("Pre y post operatorio",      CROSS,CHECK,CHECK),
    ("Gestion de quirofanos",      CROSS,CHECK,CHECK),
    ("HOSPITALIZACION",            None,None,None),
    ("Gestion de camas / censos",  CROSS,CROSS,CHECK),
    ("Notas de enfermeria IA",     CROSS,CROSS,CHECK),
    ("Evoluciones medicas 24/7",   CROSS,CROSS,CHECK),
    ("Epicrisis automatica IA",    CROSS,CROSS,CHECK),
    ("FACTURACION Y NORMATIVIDAD", None,None,None),
    ("RIPS automatico completo",   CROSS,CHECK,CHECK),
    ("Factura electronica DIAN",   SOON,SOON,SOON),
    ("Interoperabilidad HL7 FHIR", SOON,SOON,SOON),
    ("SOPORTE Y CAPACIDAD",        None,None,None),
    ("Usuarios concurrentes",      "3","15","Ilim."),
    ("Almacenamiento",             "10 GB","50 GB","500 GB"),
    ("Soporte tecnico",            "5x8","7x12","24/7"),
    ("Capacitacion inicial",       CHECK,CHECK,CHECK),
    ("Actualizaciones sistema",    CHECK,CHECK,CHECK),
]
for feat, a, b, c in sections_feats:
    row = feat_t.add_row()
    if a is None:  # Section header
        mc = row.cells[0].merge(row.cells[3])
        cell_bg(mc, NAVY2)
        rn(mc.paragraphs[0], feat, bold=True, color=GOLD, size=9)
    else:
        vals = [feat, a, b, c]
        bgs  = [LGRAY, WHITE, GBG, WHITE]
        for ci, (v, bg) in enumerate(zip(vals, bgs)):
            cl = row.cells[ci]; cell_bg(cl, bg); cell_valign(cl)
            pp = cl.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            if ci == 0:
                rn(pp, "  " + str(v), color=NAVY, size=9.5)
            elif v == CHECK:
                rn(pp, CHECK, bold=True, color=EMERALD, size=11)
            elif v == CROSS:
                rn(pp, CROSS, color=MGRAY, size=10)
            elif v == SOON:
                rn(pp, "Proximo", color=GOLD, size=8.5, italic=True)
            else:
                rn(pp, str(v), color=DGRAY, size=9)

for row in feat_t.rows:
    row.cells[0].width = Cm(7)
    for ci in range(1,4): row.cells[ci].width = Cm(3)

doc.add_page_break()

# ─ Add-ons ────────────────────────────────────────────────────────────────────
section_label(doc, "Modulos Add-On — Complementa tu Plataforma")
at = doc.add_table(rows=3, cols=3); tbl_borders(at, color="C9A84C", sz=4)
tbl_header_row(at, 0, ["MODULO ADD-ON", "PRECIO / MES + IVA", "QUE INCLUYE"], [NAVY,NAVY,NAVY], size=9)
addons = [
    ("Facturacion RIPS + DIAN",
     "$890.000",
     "RIPS completo (AC, AF, AM, AH, AU, AT, AP), Factura Electronica DIAN, cuentas de cobro EPS/ARL, glosas automaticas y conciliacion de cartera."),
    ("Interoperabilidad HL7 FHIR R4",
     "A cotizar",
     "HL7 FHIR R4 estandar internacional, conectividad MIPRES/SIVIGILA, API REST para laboratorios e imagenes, Historia Clinica Unificada RNEC, Telemedicina (Proximo)."),
]
for ri, (name, price, desc) in enumerate(addons):
    r = ri + 1
    bgs = [GBG, WHITE]
    for ci, (v, co, sz, bd) in enumerate([
        (name, NAVY, 10.5, True),(price, GDARK, 13, True),(desc, DGRAY, 9.5, False)
    ]):
        cl = at.cell(r, ci); cell_bg(cl, bgs[ri]); cell_valign(cl)
        pp = cl.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci == 1 else WD_ALIGN_PARAGRAPH.LEFT
        rn(pp, v, bold=bd, color=co, size=sz)
for row in at.rows:
    row.cells[0].width = Cm(4.5); row.cells[1].width = Cm(3.5); row.cells[2].width = Cm(9)

doc.add_page_break()

# ── 05 FLUJO CLINICO ──────────────────────────────────────────────────────────
section_label(doc, "05  Flujo Clinico Automatizado")
heading(doc, "De la cita al documento en 6 pasos automaticos", "6 pasos")
body(doc, "Todo el ciclo medico-administrativo ocurre dentro de SARAI, sin saltar de sistema en sistema.")

flow = [
    ("PASO 1","Agenda Inteligente",
     "El medico agenda la cita. El paciente recibe confirmacion por email automaticamente en segundos."),
    ("PASO 2","Check-in y Triaje",
     "El paciente confirma asistencia. El sistema actualiza el estado en tiempo real."),
    ("PASO 3","Consulta con Dictado IA",
     "El medico dicta por voz. SARAI transcribe, estructura y archiva la historia clinica."),
    ("PASO 4","Cierre de Consulta",
     "El medico marca la cita como completada. El flujo siguiente se activa automaticamente."),
    ("PASO 5","Historia Clinica Entregada",
     "La HC queda firmada digitalmente. Cumple Resolucion 1995/1999 del Ministerio de Salud."),
    ("PASO 6","Cotizacion y Facturacion",
     "El sistema genera cotizacion con totales automaticos y archivos RIPS listos para envio."),
]
ft = doc.add_table(rows=len(flow), cols=3); tbl_borders(ft, color="C9A84C", sz=4)
step_colors = [CYAN, NAVY, EMERALD, VIOLET, ROSE, GDARK]
for ri, (step, title, desc) in enumerate(flow):
    cs = ft.cell(ri, 0); ct_ = ft.cell(ri, 1); cd = ft.cell(ri, 2)
    cell_bg(cs, step_colors[ri])
    cell_bg(ct_, NAVY if ri % 2 == 0 else NAVY2)
    cell_bg(cd, LGRAY if ri % 2 == 0 else WHITE)
    cell_valign(cs); cell_valign(ct_); cell_valign(cd)
    p1 = cs.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(p1, step, bold=True, color=WHITE, size=9)
    rn(ct_.paragraphs[0], title, bold=True, color=GOLD, size=10)
    rn(cd.paragraphs[0], desc, color=DGRAY, size=9.5)
    cs.width = Cm(2); ct_.width = Cm(5); cd.width = Cm(10)

doc.add_page_break()

# ── 06 IMPLEMENTACION ────────────────────────────────────────────────────────
section_label(doc, "06  Plan de Implementacion")
heading(doc, "Operativo en 4 semanas garantizadas", "4 semanas")
body(doc, "Proceso estructurado y acompanado por nuestro equipo de expertos tecnicos y funcionales.")

tl = doc.add_table(rows=4, cols=3); tbl_borders(tl, color="C9A84C", sz=4)
timeline = [
    ("SEMANA 1", "Kickoff y Configuracion Inicial",
     "Reunion de inicio, levantamiento de requerimientos, configuracion de entorno cloud, parametrizacion de la clinica, carga de especialidades, CUPS y usuarios."),
    ("SEMANA 2", "Migracion de Datos y Capacitacion",
     "Migracion de pacientes e historial existente. Capacitacion presencial/virtual a medicos, enfermeria y administracion. Configuracion de agenda y plantillas."),
    ("SEMANA 3", "Prueba Piloto Supervisada",
     "Operacion en paralelo durante 5 dias habiles con acompanamiento en sitio. Ajustes finos, resolucion de dudas y validacion de flujos clinicos completos."),
    ("SEMANA 4  GO LIVE", "Produccion Total",
     "La clinica opera 100% en SARAI. Monitoreo activo las primeras 2 semanas post-lanzamiento. Informes de desempeno y optimizacion continua."),
]
tl_bgs = [GBG, WHITE, GBG, "E8FFF0"]
tl_colors = [NAVY, NAVY, NAVY, EMERALD]
for ri, (week, title, desc) in enumerate(timeline):
    cw = tl.cell(ri, 0); cti = tl.cell(ri, 1); cd = tl.cell(ri, 2)
    cell_bg(cw, NAVY if ri < 3 else EMERALD)
    cell_bg(cti, tl_bgs[ri]); cell_bg(cd, tl_bgs[ri])
    cell_valign(cw); cell_valign(cti); cell_valign(cd)
    p1 = cw.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(p1, week, bold=True, color=WHITE, size=9.5)
    rn(cti.paragraphs[0], title, bold=True, color=tl_colors[ri], size=11)
    rn(cd.paragraphs[0], desc, color=DGRAY, size=9.5)
    cw.width = Cm(2.8); cti.width = Cm(4.7); cd.width = Cm(9.7)

# ── 07 SLA Y GARANTIAS ────────────────────────────────────────────────────────
section_label(doc, "07  Garantia y Niveles de Servicio (SLA)")
heading(doc, "Compromisos inquebrantables con su operacion", "inquebrantables")

sla_t = doc.add_table(rows=2, cols=3); tbl_borders(sla_t, color="C9A84C", sz=4)
slas = [
    ("99.9%",     "Uptime garantizado en contrato SLA"),
    ("< 4 horas", "Respuesta a incidentes criticos"),
    ("Diario",    "Backup automatico con retencion 90 dias"),
    ("Gratis",    "Actualizaciones siempre incluidas"),
    ("Ilimitada", "Capacitacion para nuevos usuarios"),
    ("100%",      "Conformidad con normativa colombiana"),
]
for idx, (val, desc) in enumerate(slas):
    ri = idx // 3; ci = idx % 3
    cl = sla_t.cell(ri, ci)
    cell_bg(cl, GBG if ci != 1 else WHITE)
    cell_valign(cl)
    pp = cl.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rn(pp, val + "\n", bold=True, color=GDARK, size=16)
    rn(pp, desc, color=DGRAY, size=9)
    cl.width = Cm(5.7)

doc.add_paragraph()
norm = doc.add_paragraph()
rn(norm, "MARCO NORMATIVO CUMPLIDO:  ", bold=True, color=NAVY, size=10)
rn(norm, "Res. 1995/1999 (HC)  |  Res. 3374/2000 (RIPS)  |  Res. 5521/2013 (CUPS)  |  Ley 1581/2012 (Habeas Data)  |  DIAN FE  |  HL7 FHIR R4 (proximo)",
   color=DGRAY, size=9.5)

doc.add_page_break()

# ── 08 CONTACTO / CTA ────────────────────────────────────────────────────────
section_label(doc, "08  Proximos Pasos")
heading(doc, "Listo para transformar su clinica?", "transformar")
body(doc, "Agendemos una demostracion en vivo de 45 minutos donde vera SARAI funcionando con datos reales de su especialidad. Sin compromiso.", size=11)

cta = doc.add_table(rows=1, cols=2); tbl_borders(cta, color="C9A84C", sz=6)
lc = cta.cell(0,0); rc = cta.cell(0,1)
cell_bg(lc, NAVY); cell_bg(rc, GBG)

rn(lc.paragraphs[0], "SOLICITAR DEMO GRATUITA\n", bold=True, color=GOLD, size=13)
lc.add_paragraph(); rn(lc.add_paragraph(), "Email:   comercial@sarai.com.co", color=WHITE, size=10.5)
rn(lc.add_paragraph(), "Telefono:  +57 300 000 0000", color=WHITE, size=10.5)
rn(lc.add_paragraph(), "Web:      www.sarai.com.co", color=WHITE, size=10.5)

rn(rc.paragraphs[0], "GARANTIA DE SATISFACCION\n", bold=True, color=NAVY, size=12)
for g in ["Respuesta en menos de 2 horas habiles",
          "Sin letra pequena ni compromisos ocultos",
          "Cancelacion sin penalidad — primeros 30 dias",
          "Implementacion en 4 semanas garantizadas",
          "Demo con datos reales de su especialidad"]:
    p = rc.add_paragraph("  " + CHECK + "  " + g)
    p.paragraph_format.space_after = Pt(3)
    p.runs[0].font.size = Pt(9.5); p.runs[0].font.color.rgb = rgb(DGRAY)

lc.width = Cm(8); rc.width = Cm(8)

doc.add_paragraph()
divider(doc)
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.paragraph_format.space_before = Pt(12)
rn(footer, "SARAI Technologies  |  Documento Confidencial  |  Mayo 2026  |  Todos los derechos reservados",
   color=GOLD, size=8.5)

# ── GUARDAR ───────────────────────────────────────────────────────────────────
doc.save(OUTPUT)
print(f"\n  Documento creado exitosamente: {OUTPUT}\n")
